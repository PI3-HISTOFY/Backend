import cv2
import pytesseract
from pytesseract import Output
from docx import Document
from docx.shared import Pt
import numpy as np

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def procesar_imagen(ruta_imagen):
    img = cv2.imread(ruta_imagen)

    # Escalar para mejorar OCR
    altura = 2000
    ratio = altura / img.shape[0]
    img = cv2.resize(img, (int(img.shape[1] * ratio), altura))

    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    gray = cv2.medianBlur(gray, 3)
    thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                                   cv2.THRESH_BINARY_INV, 15, 10)

    # Detectar líneas de tabla
    horizontal = thresh.copy()
    vertical = thresh.copy()

    h_size = int(horizontal.shape[1] / 30)
    h_structure = cv2.getStructuringElement(cv2.MORPH_RECT, (h_size, 1))
    horizontal = cv2.erode(horizontal, h_structure)
    horizontal = cv2.dilate(horizontal, h_structure)

    v_size = int(vertical.shape[0] / 30)
    v_structure = cv2.getStructuringElement(cv2.MORPH_RECT, (1, v_size))
    vertical = cv2.erode(vertical, v_structure)
    vertical = cv2.dilate(vertical, v_structure)

    table_mask = cv2.add(horizontal, vertical)

    contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    contours = sorted(contours, key=lambda c: cv2.boundingRect(c)[1])

    doc = Document()
    texto_final = ""

    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        if w > 50 and h > 20:
            table_roi = img[y:y+h, x:x+w]
            
            gray_cell = cv2.cvtColor(table_roi, cv2.COLOR_BGR2GRAY)
            _, cell_thresh = cv2.threshold(gray_cell, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)

            cell_contours, _ = cv2.findContours(cell_thresh, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
            cells_sorted = sorted(cell_contours, key=lambda c: (cv2.boundingRect(c)[1], cv2.boundingRect(c)[0]))

            rows = []
            current_row_y = -1
            current_row = []
            
            for ccell in cells_sorted:
                cx, cy, cw, ch = cv2.boundingRect(ccell)
                if cw < 15 or ch < 15:
                    continue
                if current_row_y == -1:
                    current_row_y = cy
                if abs(cy - current_row_y) > 15:
                    rows.append(current_row)
                    current_row = []
                    current_row_y = cy
                cell_img = table_roi[cy:cy+ch, cx:cx+cw]
                cell_text = pytesseract.image_to_string(cell_img, lang="spa").strip()
                current_row.append(cell_text)
            
            if current_row:
                rows.append(current_row)

            table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    table.cell(r_idx, c_idx).text = cell_text

    # OCR de texto narrativo
    data = pytesseract.image_to_data(gray, lang="spa", output_type=Output.DICT)
    current_line_y = -1
    line_buffer = []

    def flush_line():
        nonlocal line_buffer, texto_final
        if line_buffer:
            line_text = " ".join(line_buffer)
            texto_final += line_text + "\n"
            p = doc.add_paragraph()
            run = p.add_run(line_text)
            run.font.size = Pt(11)
            line_buffer = []

    for i in range(len(data['text'])):
        if data['text'][i].strip() != "":
            y = data['top'][i]
            if current_line_y == -1:
                current_line_y = y
            if abs(y - current_line_y) > 15:
                flush_line()
                current_line_y = y
            line_buffer.append(data['text'][i])
    flush_line()

    # Guardar resultado
    output_docx = "resultado_hibrido3.docx"
    doc.save(output_docx)

    return texto_final, output_docx
